import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_compilation():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Colección de Poesía de Mujeres', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    p = doc.add_paragraph('Antología de Poetas en Lengua Española y Otros Países')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(24)
    
    doc.add_page_break()
    
    # Presentation
    doc.add_heading('Presentación', level=1)
    doc.add_paragraph(
        "Esta compilación reúne la obra y biografía de destacadas poetisas de más de 20 países. "
        "El objetivo es preservar y difundir el legado literario de mujeres que han marcado "
        "la historia de la poesía en español y otras latitudes, asegurando una representación "
        "equitativa y respetando el estatus de dominio público de las obras incluidas."
    )
    
    doc.add_page_break()
    
    # Index placeholder (Note: python-docx doesn't easily generate clickable TOCs without complex XML)
    # We will just list the sections for now.
    doc.add_heading('Índice de Contenidos', level=1)
    
    root_dir = r"c:\Users\OscarSson\Videos\GitHub\PoesiaMujeres"
    countries = sorted([d for d in os.listdir(root_dir) if os.path.isdir(os.path.join(root_dir, d)) and d != '.gemini'])
    
    # Ensure Otros_Paises is at the end
    if 'Otros_Paises' in countries:
        countries.remove('Otros_Paises')
        countries.append('Otros_Paises')
    
    # First pass for Index
    for country in countries:
        doc.add_paragraph(country.replace('_', ' '), style='List Bullet')
        
    doc.add_page_break()
    
    # Content
    for country in countries:
        country_path = os.path.join(root_dir, country)
        doc.add_heading(country.replace('_', ' '), level=1)
        
        files = sorted([f for f in os.listdir(country_path) if f.endswith('.md')])
        for file in files:
            file_path = os.path.join(country_path, file)
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Simple Markdown parsing to Docx
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    p = doc.add_heading(line[2:].strip(), level=2)
                elif line.startswith('## '):
                    p = doc.add_heading(line[3:].strip(), level=3)
                elif line.startswith('### '):
                    p = doc.add_heading(line[4:].strip(), level=4)
                elif line.strip():
                    # Handle bold
                    line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    doc.add_paragraph(line.strip())
            
            doc.add_paragraph('\n') # Spacer between poets
            
    output_path = os.path.join(root_dir, 'Compilacion_Poesia_Mujeres.docx')
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    create_compilation()
