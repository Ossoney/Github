#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Ajustador de tamaño de fuente en tablas de documentos Word (.docx)
Versión: 1.0
Autor: Asistente IA
Descripción: Reduce a 10pt cualquier texto en tablas con tamaño > 10pt,
             manteniendo el resto del documento intacto.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def obtener_tamano_fuente(run):
    """
    Obtiene el tamaño de fuente de un run (fragmento de texto).
    Retorna el tamaño en puntos o None si no está definido.
    """
    try:
        # Intentar obtener tamaño directamente del run
        if run.font.size:
            return run.font.size.pt
    except:
        pass
    
    try:
        # Buscar en el XML del run (para tamaños heredados o implícitos)
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                sz_val = sz.get(qn('w:val'))
                if sz_val:
                    # El tamaño en XML está en half-points (1/2 punto)
                    return float(sz_val) / 2
    except:
        pass
    
    return None


def establecer_tamano_fuente(run, tamano_pt):
    """
    Establece el tamaño de fuente de un run al valor especificado.
    Si el run no tenía tamaño definido, se crea la propiedad.
    """
    run.font.size = Pt(tamano_pt)
    
    # También actualizar el XML directamente para asegurar consistencia
    rPr = run._element.get_or_add_rPr()
    
    # Eliminar cualquier tamaño existente en XML
    for sz in rPr.findall(qn('w:sz')):
        rPr.remove(sz)
    for szCs in rPr.findall(qn('w:szCs')):
        rPr.remove(szCs)
    
    # Crear nuevo elemento de tamaño (en half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(tamano_pt * 2)))
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(tamano_pt * 2)))
    rPr.append(szCs)


def procesar_tablas_en_documento(doc, max_tamano=10, dry_run=False):
    """
    Procesa todas las tablas en el documento.
    Si dry_run es True, solo cuenta sin modificar.
    Retorna estadísticas de modificaciones.
    """
    stats = {
        'tablas_procesadas': 0,
        'runs_modificados': 0,
        'runs_ya_ok': 0,
        'runs_sin_tamano_definido': 0
    }
    
    # Recorrer todas las tablas
    for tabla_idx, tabla in enumerate(doc.tables):
        tabla_modificada = False
        
        # Recorrer todas las filas
        for fila in tabla.rows:
            # Recorrer todas las celdas
            for celda in fila.cells:
                # Cada celda puede tener múltiples párrafos
                for parrafo in celda.paragraphs:
                    # Cada párrafo puede tener múltiples runs
                    for run in parrafo.runs:
                        # Obtener tamaño actual
                        tamano_actual = obtener_tamano_fuente(run)
                        
                        if tamano_actual is None:
                            # No tiene tamaño definido (usa el predeterminado del estilo)
                            stats['runs_sin_tamano_definido'] += 1
                            continue
                        
                        if tamano_actual > max_tamano:
                            # Necesita reducción
                            if not dry_run:
                                establecer_tamano_fuente(run, max_tamano)
                            stats['runs_modificados'] += 1
                            tabla_modificada = True
                        else:
                            # Ya está en tamaño correcto
                            stats['runs_ya_ok'] += 1
        
        if tabla_modificada:
            stats['tablas_procesadas'] += 1
    
    return stats


def main():
    """
    Función principal: procesa argumentos, abre documento, procesa y guarda.
    """
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Reduce a 10pt el texto en tablas de documentos Word',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  %(prog)s documento.docx
  %(prog)s documento.docx --salida documento_ajustado.docx
  %(prog)s documento.docx --maximo 12 --dry-run
  %(prog)s *.docx --directorio ./procesados
        """
    )
    
    parser.add_argument(
        'archivos',
        nargs='+',
        help='Archivo(s) .docx a procesar (pueden ser comodines como *.docx)'
    )
    
    parser.add_argument(
        '--maximo', '-m',
        type=float,
        default=10.0,
        help='Tamaño máximo permitido en tablas (defecto: 10.0)'
    )
    
    parser.add_argument(
        '--salida', '-o',
        help='Nombre del archivo de salida (solo si se procesa un único archivo)'
    )
    
    parser.add_argument(
        '--directorio', '-d',
        help='Directorio de salida para archivos procesados'
    )
    
    parser.add_argument(
        '--dry-run', '-n',
        action='store_true',
        help='Solo simular, no guardar cambios (útil para probar)'
    )
    
    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Mostrar información detallada por cada tabla'
    )
    
    args = parser.parse_args()
    
    # Expandir archivos si hay comodines
    archivos_procesar = []
    for patron in args.archivos:
        if '*' in patron or '?' in patron:
            from glob import glob
            archivos_procesar.extend(glob(patron))
        else:
            archivos_procesar.append(patron)
    
    # Validar que haya al menos un archivo
    if not archivos_procesar:
        print("❌ No se encontraron archivos para procesar")
        sys.exit(1)
    
    # Validar que los archivos existan
    archivos_existentes = []
    for archivo in archivos_procesar:
        if Path(archivo).exists():
            archivos_existentes.append(archivo)
        else:
            print(f"⚠️  Archivo no encontrado: {archivo}")
    
    if not archivos_existentes:
        print("❌ Ningún archivo válido encontrado")
        sys.exit(1)
    
    # Si hay múltiples archivos, no se puede usar --salida
    if len(archivos_existentes) > 1 and args.salida:
        print("❌ No se puede usar --salida con múltiples archivos")
        print("   Use --directorio para especificar carpeta de destino")
        sys.exit(1)
    
    # Crear directorio de salida si se especificó
    if args.directorio:
        Path(args.directorio).mkdir(parents=True, exist_ok=True)
    
    # Procesar cada archivo
    total_archivos = len(archivos_existentes)
    
    print(f"\n📄 Procesando {total_archivos} archivo(s)...")
    print(f"⚙️  Tamaño máximo permitido: {args.maximo}pt")
    print(f"🔍 Modo simulación: {'SÍ' if args.dry_run else 'NO'}\n")
    
    estadisticas_totales = {
        'archivos_procesados': 0,
        'archivos_modificados': 0,
        'total_tablas_procesadas': 0,
        'total_runs_modificados': 0
    }
    
    for idx, archivo in enumerate(archivos_existentes, 1):
        print(f"[{idx}/{total_archivos}] Procesando: {archivo}")
        
        try:
            # Abrir documento
            doc = Document(archivo)
            
            # Procesar tablas
            stats = procesar_tablas_en_documento(doc, args.maximo, args.dry_run)
            
            # Mostrar resultados
            print(f"   📊 Tablas con cambios: {stats['tablas_procesadas']}")
            if args.verbose:
                print(f"      - Runs modificados: {stats['runs_modificados']}")
                print(f"      - Runs ya en tamaño correcto: {stats['runs_ya_ok']}")
                print(f"      - Runs sin tamaño definido: {stats['runs_sin_tamano_definido']}")
            
            # Guardar si hay cambios y no es dry-run
            if stats['runs_modificados'] > 0:
                if not args.dry_run:
                    # Determinar nombre de salida
                    if args.salida:
                        nombre_salida = args.salida
                    elif args.directorio:
                        nombre_base = Path(archivo).name
                        nombre_salida = str(Path(args.directorio) / nombre_base)
                    else:
                        # Por defecto: nombre_original_ajustado.docx
                        ruta = Path(archivo)
                        nombre_salida = str(ruta.parent / f"{ruta.stem}_ajustado{ruta.suffix}")
                    
                    # Guardar documento
                    doc.save(nombre_salida)
                    print(f"   💾 Guardado en: {nombre_salida}")
                    estadisticas_totales['archivos_modificados'] += 1
                else:
                    print("   🔍 [SIMULACIÓN] No se guardaron cambios")
            
            estadisticas_totales['archivos_procesados'] += 1
            estadisticas_totales['total_tablas_procesadas'] += stats['tablas_procesadas']
            estadisticas_totales['total_runs_modificados'] += stats['runs_modificados']
            
        except Exception as e:
            print(f"   ❌ Error procesando {archivo}: {e}")
        
        print()  # Línea en blanco entre archivos
    
    # Resumen final
    print("=" * 50)
    print("📋 RESUMEN FINAL")
    print(f"   Archivos procesados: {estadisticas_totales['archivos_procesados']}")
    print(f"   Archivos modificados: {estadisticas_totales['archivos_modificados']}")
    print(f"   Tablas procesadas: {estadisticas_totales['total_tablas_procesadas']}")
    print(f"   Fragmentos de texto modificados: {estadisticas_totales['total_runs_modificados']}")
    
    if args.dry_run:
        print("\n⚠️  Este fue un SIMULACIÓN (dry-run). Ningún archivo fue modificado realmente.")
        print("   Ejecuta sin --dry-run para aplicar los cambios.")


if __name__ == "__main__":
    main()