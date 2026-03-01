"""
Script para llenar archivos de poetas con sus 10 mejores POEMAS REALES de amor.
Modo: SOLO poemas reales documentados. Si no hay 10, se indica cuántos hay.
Colecciones:
  - poetas-en-frances/** (poemas en francés + traducción al español)
  - poetas_chinos_libro    (chino clásico + español)
  - poetisas_chinas_libro  (chino clásico + español)
  - poetas_japoneses_libro (japonés + español)

Uso:
  $env:GEMINI_API_KEY="TU_CLAVE"; python generate_poems.py
"""

import os, sys, re, json, time, glob, subprocess

# ─── SDK ────────────────────────────────────────────────────────────────────
try:
    from google import genai as gai
    from google.genai import types as gai_types
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "google-genai"])
    from google import genai as gai
    from google.genai import types as gai_types

try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
    from docx import Document

# ─── Configuración ──────────────────────────────────────────────────────────
API_KEY = os.environ.get("GEMINI_API_KEY", "")
if not API_KEY:
    print("ERROR: falta la variable de entorno GEMINI_API_KEY")
    sys.exit(1)

BASE = r"c:\Users\OscarSson\Videos\GitHub\aaaa"
GENAI_MODEL = "gemini-2.0-pro-exp-02-05"

COLLECTIONS = [
    {
        "label": "poetas-en-frances",
        "md_glob": os.path.join(BASE, "poetas-en-frances", "**", "*_libro", "*.md"),
        "lang_original": "francés",
        "lang_label_original": "Francés (Original)",
        "lang_label_traduccion": "Español (Traducción)",
    },
    {
        "label": "poetas_chinos_libro",
        "md_glob": os.path.join(BASE, "poetas_chinos_libro", "*.md"),
        "lang_original": "chino clásico",
        "lang_label_original": "Original (Chino)",
        "lang_label_traduccion": "Traducción (Español)",
    },
    {
        "label": "poetisas_chinas_libro",
        "md_glob": os.path.join(BASE, "poetisas_chinas_libro", "*.md"),
        "lang_original": "chino clásico",
        "lang_label_original": "Original (Chino)",
        "lang_label_traduccion": "Traducción (Español)",
    },
    {
        "label": "poetas_japoneses_libro",
        "md_glob": os.path.join(BASE, "poetas_japoneses_libro", "*.md"),
        "lang_original": "japonés",
        "lang_label_original": "Original (Japonés)",
        "lang_label_traduccion": "Traducción (Español)",
    },
    {
        "label": "poetisas_japonesas_libro",
        "md_glob": os.path.join(BASE, "poetisas_japonesas_libro", "*.md"),
        "lang_original": "japonés",
        "lang_label_original": "Original (Japonés)",
        "lang_label_traduccion": "Traducción (Español)",
    },
    {
        "label": "poetisas_italianas_libro",
        "md_glob": os.path.join(BASE, "poetisas_italianas_libro", "*.md"),
        "lang_original": "italiano",
        "lang_label_original": "Italiano",
        "lang_label_traduccion": "Castellano",
    },
]

PLACEHOLDER_HINTS = [
    "Te voilà revenu",
    "Sin sombras te amo",
    "Unidos por los mismos espíritus",
    "love you without shadows",
    "sin sombras te amo",
]

# ─── Detección ───────────────────────────────────────────────────────────────

def is_incomplete(content: str) -> bool:
    for ph in PLACEHOLDER_HINTS:
        if ph.lower() in content.lower():
            return True
    rows = re.findall(r"\|([^|\n]+)\|[^|\n]+\|", content)
    skip = {"---", "original (chino)", "original (japonés)", "francés (original)",
            "español (traducción)", "traducción (español)"}
    cells = [r.strip() for r in rows if r.strip().lower() not in skip and len(r.strip()) > 5]
    if len(cells) >= 10 and len(set(cells[:10])) == 1:
        return True  # mismo poema repetido 10 veces
    return False

# ─── Utilidades ──────────────────────────────────────────────────────────────

def poet_name_from_path(filepath: str) -> str:
    return os.path.basename(filepath).replace(".md", "").replace("_", " ")

def extract_bio(content: str) -> str:
    m = re.search(r"## Biografía Sentimental\s*\n(.+?)(?=\n##|\Z)", content, re.DOTALL)
    if m:
        bio = m.group(1).strip()
        if bio and not any(ph.lower() in bio.lower() for ph in PLACEHOLDER_HINTS):
            return bio[:800]
    return ""

def find_docx_dir(md_path: str) -> str | None:
    docx_dir = os.path.dirname(md_path).replace("_libro", "_docx")
    return docx_dir if os.path.isdir(docx_dir) else None

# ─── Prompt y API ────────────────────────────────────────────────────────────

def build_prompt(poet_name: str, lang: str, bio_hint: str) -> str:
    bio_section = f"\nBiografía conocida:\n{bio_hint}\n" if bio_hint else ""
    return f"""Eres un experto en literatura de {lang}. Necesito los mejores poemas de amor REALES y DOCUMENTADOS del poeta **{poet_name}**.

Responde ÚNICAMENTE con JSON válido. Sin texto antes ni después.

{{
  "bio": "<3-5 oraciones REALES sobre la vida amorosa/sentimental de {poet_name}, con hechos históricos concretos, en español>",
  "poems_found": <número entero: cuántos poemas de amor reales y documentados tiene este poeta>,
  "poems": [
    {{
      "title": "<título real y documentado del poema en {lang}>",
      "original": "<texto completo del poema en {lang}, saltos de línea = \\n>",
      "traduccion": "<traducción poética al español, saltos de línea = \\n>"
    }}
  ]
}}

REGLAS ESTRICTAS:
1. Incluye ÚNICAMENTE poemas que sean obras REALES y DOCUMENTADAS de {poet_name}.
2. NO inventes poemas. NO compongas poemas nuevos. NO atribuyas poemas ajenos.
3. Si el poeta tiene menos de 10 poemas de amor documentados, incluye SOLO los que existen realmente.
4. El campo "poems_found" debe reflejar honestamente cuántos poemas reales de amor se conocen de este autor.
5. Los poemas deben ser DISTINTOS entre sí: diferentes títulos, diferentes primeros versos.
6. Texto original en {lang} con escritura nativa (chino, japonés, etc. según aplique).
7. Si este autor NO es conocido principalmente como poeta (p.ej. es novelista, ensayista, etnógrafo...), indícalo en la bio y pon solo los poemas que realmente escribió.{bio_section}"""


def generate_for_poet(client, poet_name: str, lang: str, bio_hint: str) -> dict | None:
    prompt = build_prompt(poet_name, lang, bio_hint)
    for attempt in range(3):
        try:
            response = client.models.generate_content(
                model=GENAI_MODEL,
                contents=prompt,
                config=gai_types.GenerateContentConfig(
                    temperature=0.2,   # bajo para maximizar fidelidad factual
                    max_output_tokens=8192,
                )
            )
            text = response.text.strip()
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```\s*$", "", text)
            data = json.loads(text)
            if "poems" in data:
                return data
            print(f"    [aviso] JSON sin 'poems', reintentando ({attempt+1}/3)...")
        except json.JSONDecodeError as e:
            print(f"    [JSON error] intento {attempt+1}: {e}")
            if attempt == 0:
                print(f"    Respuesta raw: {text[:400]}")
        except Exception as e:
            print(f"    [API error] intento {attempt+1}: {e}")
            time.sleep(8)
    return None

# ─── Construcción de archivos ─────────────────────────────────────────────────

NOTE_TEMPLATE = "*(No se han encontrado {n} poemas de amor documentados para este autor. " \
                "Solo existen {found} registrados en la literatura conocida.)*"

def build_md(poet_name: str, bio: str, poems: list, poems_found: int, col: dict) -> str:
    lo = col["lang_label_original"]
    lt = col["lang_label_traduccion"]
    lines = [f"# {poet_name}", "", "## Biografía Sentimental", "", bio, "", "## Poemas de Amor", ""]

    for i, p in enumerate(poems, 1):
        orig = p.get("original", "").replace("\\n", "\n")
        trad = p.get("traduccion", "").replace("\\n", "\n")
        lines += [
            f"### {i}. {p.get('title', f'Poema {i}')}",
            "",
            f"| {lo} | {lt} |",
            "|---|---|",
            f"| {orig.replace(chr(10), '<br>')} | {trad.replace(chr(10), '<br>')} |",
            "",
        ]

    if len(poems) < 10:
        lines += [
            f"---",
            "",
            NOTE_TEMPLATE.format(n=10, found=len(poems)),
            f"*(Poemas de amor reales documentados: {poems_found if poems_found else len(poems)} en total)*",
            "",
        ]
    return "\n".join(lines)


def build_docx(poet_name: str, bio: str, poems: list, poems_found: int, out_path: str, col: dict):
    lo = col["lang_label_original"]
    lt = col["lang_label_traduccion"]
    doc = Document()
    doc.add_heading(poet_name, 0)
    doc.add_heading("Biografía Sentimental", level=1)
    doc.add_paragraph(bio)
    doc.add_heading("Poemas de Amor", level=1)
    for i, p in enumerate(poems, 1):
        doc.add_heading(f"{i}. {p.get('title', f'Poema {i}')}", level=2)
        tbl = doc.add_table(rows=2, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = lo
        tbl.rows[0].cells[1].text = lt
        tbl.rows[1].cells[0].text = p.get("original", "").replace("\\n", "\n")
        tbl.rows[1].cells[1].text = p.get("traduccion", "").replace("\\n", "\n")
        doc.add_paragraph()
    if len(poems) < 10:
        doc.add_paragraph(NOTE_TEMPLATE.format(n=10, found=len(poems)))
        doc.add_paragraph(f"Poemas de amor reales documentados: {poems_found if poems_found else len(poems)} en total")
    doc.save(out_path)

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    client = gai.Client(api_key=API_KEY)
    total_fixed = 0
    total_errors = 0
    poets_with_few = []

    for col in COLLECTIONS:
        md_files = glob.glob(col["md_glob"], recursive=True)
        to_fix = []
        for f in md_files:
            if "NO_DERECHOS" in f:
                continue
            try:
                content = open(f, encoding="utf-8").read()
                if is_incomplete(content):
                    to_fix.append((f, content))
            except Exception as e:
                print(f"  [read error] {f}: {e}")

        print(f"\n{'='*65}")
        print(f"Colección: {col['label']}  →  {len(to_fix)} archivos a recuperar")
        print(f"{'='*65}")

        for md_path, content in to_fix:
            poet = poet_name_from_path(md_path)
            bio_hint = extract_bio(content)
            print(f"\n  ▶ {poet} ...")

            data = generate_for_poet(client, poet, col["lang_original"], bio_hint)
            if not data:
                print(f"  ✗ FALLO: {poet}")
                total_errors += 1
                continue

            bio = data.get("bio", bio_hint or "")
            poems = data.get("poems", [])
            poems_found = data.get("poems_found", len(poems))

            if len(poems) < 10:
                poets_with_few.append(f"{poet} ({len(poems)}/10 poemas reales)")

            # Escribir .md
            open(md_path, "w", encoding="utf-8").write(
                build_md(poet, bio, poems, poems_found, col)
            )

            # Escribir .docx
            docx_dir = find_docx_dir(md_path)
            if docx_dir:
                docx_file = os.path.join(docx_dir, os.path.basename(md_path).replace(".md", ".docx"))
                try:
                    build_docx(poet, bio, poems, poems_found, docx_file, col)
                except Exception as e:
                    print(f"  [docx error] {e}")

            total_fixed += 1
            status = "✓" if len(poems) >= 10 else f"⚠ solo {len(poems)} poemas reales"
            print(f"  {status}  {poet}")
            time.sleep(2)

    print(f"\n{'='*65}")
    print(f"RESUMEN: {total_fixed} procesados, {total_errors} errores")
    if poets_with_few:
        print(f"\nPoetas con menos de 10 poemas reales documentados ({len(poets_with_few)}):")
        for p in poets_with_few:
            print(f"  - {p}")
    print(f"{'='*65}")


if __name__ == "__main__":
    main()
