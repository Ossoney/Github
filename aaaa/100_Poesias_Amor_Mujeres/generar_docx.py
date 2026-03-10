import os
import re
import urllib.request
from docx import Document
from docx.shared import Inches

base_path = "/home/osso/Descargas/aaaa/100_Poesias_Amor_Mujeres/poemas/"
output_path = "/home/osso/Descargas/aaaa/100_Poesias_Amor_Mujeres/100_Poesias_Amor_Mujeres.docx"

doc = Document()
doc.add_heading('Las 100 Mejores Poesías de Amor Escritas por Mujeres', 0)

files = sorted(os.listdir(base_path))

image_pattern = re.compile(r'!\[.*?\]\((.*?)\)')
title_pattern = re.compile(r'^#\s+(.*)')
subtitle_pattern = re.compile(r'^##\s+(.*)')

for filename in files:
    if not filename.endswith(".md"):
        continue
    
    filepath = os.path.join(base_path, filename)
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    in_table = False
    
    for line in lines:
        title_match = title_pattern.match(line)
        if title_match:
            doc.add_heading(title_match.group(1), level=1)
            continue
            
        subtitle_match = subtitle_pattern.match(line)
        if subtitle_match:
            doc.add_heading(subtitle_match.group(1), level=2)
            continue
            
        image_match = image_pattern.search(line)
        if image_match:
            image_url = image_match.group(1)
            temp_img = "/tmp/temp_poet_img.jpg"
            try:
                urllib.request.urlretrieve(image_url, temp_img)
                doc.add_picture(temp_img, width=Inches(2.5))
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"[Imagen no disponible: {image_url}]")
            continue
            
        if line.startswith('| Original'):
            in_table = True
            doc.add_paragraph("--- Poema Bilingüe ---")
            continue
        elif line.startswith('|:') or line.startswith('|---'):
            continue
        elif line.startswith('|') and in_table:
            parts = [p.strip() for p in line.strip('|').split('|')]
            if len(parts) >= 2:
                p = doc.add_paragraph()
                r = p.add_run(parts[0] + "  ///  " + parts[1])
                r.italic = True
            continue
        elif not line.strip() and in_table:
            in_table = False
            continue
            
        if line.strip() and not in_table:
            doc.add_paragraph(line.strip())

doc.save(output_path)
print(f"Recopilación guardada en {output_path}")
